import pandas as pd
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import customtkinter as ctk
from tkinter import filedialog, messagebox
import datetime
import os
import calendar
import sys

# ============================================================
# НАСТРОЙКА CustomTkinter
# ============================================================
ctk.set_appearance_mode("dark")      # "dark", "light", "system"
ctk.set_default_color_theme("blue")  # "blue", "dark-blue", "green"

# ============================================================
# ПУТЬ К ПАПКЕ (для портативной версии)
# ============================================================
def get_base_path():
    """Возвращает путь к папке, где находится исполняемый файл (или скрипт)"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

# ============================================================
# ПЫТАЕМСЯ ИМПОРТИРОВАТЬ БИБЛИОТЕКУ ДЛЯ PDF
# ============================================================
PDF_AVAILABLE = False
PDF_LIB = None

try:
    import pdfplumber
    PDF_AVAILABLE = True
    PDF_LIB = 'pdfplumber'
except ImportError:
    try:
        from pypdf import PdfReader
        PDF_AVAILABLE = True
        PDF_LIB = 'pypdf'
    except ImportError:
        print("⚠️ Ни одна PDF-библиотека не установлена.")

# ============================================================
# СПРАВОЧНИКИ
# ============================================================
def load_spravochnik(filename):
    """Загружает справочник из текстового файла (рядом с .exe)"""
    base_path = get_base_path()
    file_path = os.path.join(base_path, filename)
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]
        return lines
    except FileNotFoundError:
        print(f"⚠️ Файл {filename} не найден, создаём с данными по умолчанию")
        return []

def save_spravochnik(filename, data):
    """Сохраняет справочник в текстовый файл (рядом с .exe)"""
    base_path = get_base_path()
    file_path = os.path.join(base_path, filename)
    
    with open(file_path, 'w', encoding='utf-8') as f:
        for item in data:
            f.write(item + '\n')

# Загружаем справочники
RUKOVODITELI = load_spravochnik('spravochnik_rukovoditeli.txt')
PLOSHADKI = load_spravochnik('spravochnik_ploshadki.txt')

# Если справочники пустые, создаём с данными по умолчанию
if not RUKOVODITELI:
    RUKOVODITELI = [
        'Главный врач ГОБУЗ МОКМЦ|Тарбаев Е.Ю.',
        'Зам. гл. врача по медицинской части ГОБУЗ МОКМЦ|Гредягин С.С.',
        'Зам. гл. врача по терапии ГОБУЗ МОКМЦ|Колосова О.Л.'
    ]
    save_spravochnik('spravochnik_rukovoditeli.txt', RUKOVODITELI)

if not PLOSHADKI:
    PLOSHADKI = ['Ломоносова, д.18', 'Володарского, д.18', 'Перинатальный центр', 'Родильный дом']
    save_spravochnik('spravochnik_ploshadki.txt', PLOSHADKI)

# Месяцы на русском
MONTHS_RU = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
}

MONTHS_RU_NOMINATIVE = {
    1: 'Январь', 2: 'Февраль', 3: 'Март', 4: 'Апрель',
    5: 'Май', 6: 'Июнь', 7: 'Июль', 8: 'Август',
    9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'
}

MONTHS_RU_UPPER = {
    1: 'ЯНВАРЬ', 2: 'ФЕВРАЛЬ', 3: 'МАРТ', 4: 'АПРЕЛЬ',
    5: 'МАЙ', 6: 'ИЮНЬ', 7: 'ИЮЛЬ', 8: 'АВГУСТ',
    9: 'СЕНТЯБРЬ', 10: 'ОКТЯБРЬ', 11: 'НОЯБРЬ', 12: 'ДЕКАБРЬ'
}

# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ (ПАРСЕРЫ)
# ============================================================
def normalize_time(time_str):
    if not time_str or time_str == '':
        return time_str
    
    time_str = time_str.strip()
    
    match = re.match(r'(\d{1,2}):(\d{2})\s*[-—]\s*(\d{1,2}):(\d{2})', time_str)
    if match:
        h1, m1, h2, m2 = match.groups()
        return f"{int(h1):02d}:{m1}-{int(h2):02d}:{m2}"
    
    match = re.match(r'(\d{1,2})\s*[/]\s*(\d{1,2})', time_str)
    if match:
        h1, h2 = match.groups()
        return f"{int(h1):02d}-{int(h2):02d}"
    
    match = re.match(r'(\d{1,2})\s*[-—]\s*(\d{1,2})', time_str)
    if match:
        h1, h2 = match.groups()
        return f"{int(h1):02d}-{int(h2):02d}"
    
    match = re.match(r'0-(\d{1,2})', time_str)
    if match:
        h2 = match.group(1)
        return f"00-{int(h2):02d}"
    
    return time_str

def read_excel_or_csv(file_path):
    file_path = Path(file_path)
    
    if not file_path.exists():
        return None
    
    if file_path.suffix.lower() in ['.xlsx', '.xls']:
        try:
            df = pd.read_excel(file_path, header=None)
            return df
        except Exception as e:
            print(f"   ⚠️ Ошибка чтения Excel: {e}")
            return None
    
    elif file_path.suffix.lower() == '.csv':
        try:
            df = pd.read_csv(file_path, sep=';', encoding='utf-8-sig', header=None)
            return df
        except:
            try:
                df = pd.read_csv(file_path, sep=';', encoding='cp1251', header=None)
                return df
            except:
                try:
                    df = pd.read_csv(file_path, sep=';', encoding='latin-1', header=None)
                    return df
                except:
                    return None
    
    return None

def get_dept_abbr(dept_name):
    DEPT_ABBR = {
        'Ответственные по стационару': 'ОТВ',
        'Приемное отделение': 'ПО',
        'ОАР': 'ОАР',
        'Хирургия': 'ХИР',
        'Гинекология': 'ГИН',
        'Урология': 'УРО',
        'Травматология': 'ТРАВМ',
        'Травмпункт': 'ТР П',
    }
    
    if dept_name in DEPT_ABBR:
        return DEPT_ABBR[dept_name]
    
    for key, abbr in DEPT_ABBR.items():
        if key.lower() in dept_name.lower() or dept_name.lower() in key.lower():
            return abbr
    
    words = dept_name.split()
    if len(words) >= 2:
        return ''.join(w[0].upper() for w in words[:2])
    return dept_name[:4].upper()

def get_dept_order(dept_name):
    DEPARTMENT_ORDER = [
        'Ответственные по стационару',
        'Приемное отделение',
        'ОАР',
        'Хирургия',
        'Гинекология',
        'Урология',
        'Травматология',
        'Травмпункт',
    ]
    for i, dept in enumerate(DEPARTMENT_ORDER):
        if dept.lower() in dept_name.lower() or dept_name.lower() in dept.lower():
            return i
    return len(DEPARTMENT_ORDER)

def extract_day_from_date(val):
    val_str = str(val).strip()
    
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', val_str)
    if match:
        return int(match.group(3))
    
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', val_str)
    if match:
        return int(match.group(1))
    
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', val_str)
    if match:
        return int(match.group(1))
    
    match = re.search(r'\b([1-9]|[12][0-9]|3[01])\b', val_str)
    if match:
        return int(match.group(1))
    
    return None

def extract_month_from_date(val):
    val_str = str(val).strip()
    
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', val_str)
    if match:
        return int(match.group(2))
    
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', val_str)
    if match:
        return int(match.group(2))
    
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', val_str)
    if match:
        return int(match.group(2))
    
    return None

def extract_year_from_date(val):
    val_str = str(val).strip()
    
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', val_str)
    if match:
        return int(match.group(1))
    
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', val_str)
    if match:
        return int(match.group(3))
    
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', val_str)
    if match:
        return int(match.group(3))
    
    return None

def is_valid_doctor_name(name):
    if not name or not isinstance(name, str):
        return False
    
    name = name.strip()
    if not name:
        return False
    
    headers = ['Дата', 'Дежурный врач', 'Сб', 'Вс', 'Пн', 'Вт', 'Ср', 'Чт', 'Пт',
               'График', 'Утверждаю', 'Заведующий', 'Ф.И.О.', 'цех', 'отделение',
               '2026', '2025', 'Август', 'Март']
    for h in headers:
        if h in name:
            return False
    
    if re.search(r'\d{4}-\d{2}-\d{2}', name):
        return False
    
    if len(name) > 50:
        return False
    
    if not re.search(r'[а-яА-ЯёЁ]', name):
        return False
    
    return True

# ============================================================
# ПАРСЕРЫ
# ============================================================
def detect_department(df, file_path):
    file_stem = Path(file_path).stem.lower()
    
    header_text = ''
    for i in range(min(10, len(df))):
        row = df.iloc[i].values
        row_str = ' '.join(str(v) for v in row if pd.notna(v))
        header_text += row_str + ' '
    
    header_text = header_text.lower()
    
    DEPARTMENT_ORDER = [
        'Ответственные по стационару',
        'Приемное отделение',
        'ОАР',
        'Хирургия',
        'Гинекология',
        'Урология',
        'Травматология',
        'Травмпункт',
    ]
    
    for dept in DEPARTMENT_ORDER:
        dept_clean = re.sub(r'[_\s]+', ' ', dept.lower())
        clean_filename = re.sub(r'графики?_дежурных?_врачей?_', '', file_stem)
        clean_filename = re.sub(r'[_\s]+', ' ', clean_filename).strip()
        if dept_clean in clean_filename or clean_filename in dept_clean:
            return dept
    
    dept_patterns = {
        'Ответственные по стационару': ['ответственн', 'ответсвтвенн', 'стационар'],
        'Приемное отделение': ['приемн', 'приёмн', 'приемного'],
        'ОАР': ['оар', 'реаниматолог', 'анестезиолог', 'анестезия'],
        'Хирургия': ['хирург', 'хирургическ'],
        'Гинекология': ['гинеколог', 'гинекологическ'],
        'Урология': ['уролог', 'урологическ'],
        'Травматология': ['травматолог', 'травматологическ'],
        'Травмпункт': ['травмпункт', 'травм пункт'],
    }
    
    for dept, patterns in dept_patterns.items():
        for pattern in patterns:
            if pattern in header_text:
                return dept
    
    return Path(file_path).stem.replace('_', ' ')

def parse_standard_graph(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name}...")
    
    df = read_excel_or_csv(file_path)
    if df is None:
        print(f"   ⚠️ Не удалось прочитать файл")
        return {}, dept_name, None, None
    
    date_row_idx = None
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Дата' in row_str and re.search(r'\d{4}-\d{2}-\d{2}|\d{2}\.\d{2}\.\d{4}', row_str):
            date_row_idx = i
            break
    
    if date_row_idx is None:
        for i, row in df.iterrows():
            row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
            numbers = re.findall(r'\d+', row_str)
            if len(numbers) >= 10:
                date_row_idx = i
                break
    
    if date_row_idx is None:
        print(f"   ⚠️ Не найдена строка с датами")
        return {}, dept_name, None, None
    
    date_row = df.iloc[date_row_idx].values
    day_cols = []
    months_found = set()
    years_found = set()
    
    for i, val in enumerate(date_row):
        if pd.notna(val):
            day_num = extract_day_from_date(val)
            month_num = extract_month_from_date(val)
            year_num = extract_year_from_date(val)
            if month_num:
                months_found.add(month_num)
            if year_num:
                years_found.add(year_num)
            if day_num and 1 <= day_num <= 31:
                day_cols.append((i, day_num))
    
    print(f"   📅 Найдено {len(day_cols)} дней, месяц: {months_found}, год: {years_found}")
    
    header_row_idx = None
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Дежурный врач' in row_str or 'дежурный' in row_str.lower():
            header_row_idx = i
            break
    
    if header_row_idx is None:
        print(f"   ⚠️ Не найдена строка 'Дежурный врач'")
        return {}, dept_name, None, None
    
    result = {day: {} for _, day in day_cols}
    doctors_found = 0
    
    for i in range(header_row_idx + 1, len(df)):
        row = df.iloc[i].values
        if len(row) == 0 or pd.isna(row[0]) or str(row[0]).strip() == '':
            continue
        doctor = str(row[0]).strip()
        
        if not is_valid_doctor_name(doctor):
            continue
        
        if doctor == '' or doctor == 'Заведующий' or 'Заведующий' in doctor:
            break
        
        doctors_found += 1
        for col_idx, day_num in day_cols:
            if col_idx < len(row) and pd.notna(row[col_idx]):
                val = str(row[col_idx]).strip()
                if val and val not in ['', 'nan', 'None']:
                    if re.search(r'\d+[-/]\d+|\d+:\d+', val):
                        normalized = normalize_time(val)
                        result[day_num][(doctor, dept_name)] = normalized
    
    print(f"   👨‍⚕️ Найдено {doctors_found} врачей")
    
    return result, dept_name, months_found, years_found

def parse_responsible_graph(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name}...")
    
    df = read_excel_or_csv(file_path)
    if df is None:
        print(f"   ⚠️ Не удалось прочитать файл")
        return {}, dept_name, None, None
    
    date_row_idx = None
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Дата' in row_str and re.search(r'\d{4}-\d{2}-\d{2}|\d{2}\.\d{2}\.\d{4}', row_str):
            date_row_idx = i
            break
    
    if date_row_idx is None:
        for i, row in df.iterrows():
            row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
            numbers = re.findall(r'\d+', row_str)
            if len(numbers) >= 10:
                date_row_idx = i
                break
    
    if date_row_idx is None:
        print(f"   ⚠️ Не найдена строка с датами")
        return {}, dept_name, None, None
    
    date_row = df.iloc[date_row_idx].values
    day_cols = []
    months_found = set()
    years_found = set()
    
    for i, val in enumerate(date_row):
        if pd.notna(val):
            day_num = extract_day_from_date(val)
            month_num = extract_month_from_date(val)
            year_num = extract_year_from_date(val)
            if month_num:
                months_found.add(month_num)
            if year_num:
                years_found.add(year_num)
            if day_num and 1 <= day_num <= 31:
                day_cols.append((i, day_num))
    
    print(f"   📅 Найдено {len(day_cols)} дней, месяц: {months_found}, год: {years_found}")
    
    result = {day: {} for _, day in day_cols}
    doctors_found = 0
    
    for i, row in df.iterrows():
        if len(row) == 0 or pd.isna(row[0]) or str(row[0]).strip() == '':
            continue
        first_cell = str(row[0]).strip()
        
        if first_cell in ['Дата', 'Дежурный врач', 'Сб', 'Вс', 'Пн', 'Вт', 'Ср', 'Чт', 'Пт']:
            continue
        if not first_cell or first_cell == '':
            continue
        
        if not is_valid_doctor_name(first_cell):
            continue
        
        doctor = first_cell
        doctors_found += 1
        for col_idx, day_num in day_cols:
            if col_idx < len(row) and pd.notna(row[col_idx]):
                val = str(row[col_idx]).strip()
                if val and val not in ['', 'nan', 'None']:
                    if re.search(r'\d+[-/]\d+|\d+:\d+', val):
                        normalized = normalize_time(val)
                        result[day_num][(doctor, dept_name)] = normalized
    
    print(f"   👨‍⚕️ Найдено {doctors_found} врачей")
    
    return result, dept_name, months_found, years_found

def parse_oar_graph(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name}...")
    
    df = read_excel_or_csv(file_path)
    if df is None:
        print(f"   ⚠️ Не удалось прочитать файл")
        return {}, dept_name, None, None
    
    header_row_idx = None
    months_found = set()
    years_found = set()
    
    # Поиск месяца в тексте (для ОАР)
    month_names_ru = {
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4,
        'май': 5, 'июнь': 6, 'июль': 7, 'август': 8,
        'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12
    }
    
    for i in range(min(5, len(df))):
        row = df.iloc[i].values
        row_str = ' '.join(str(v) for v in row if pd.notna(v)).lower()
        for month_name, month_num in month_names_ru.items():
            if month_name in row_str:
                months_found.add(month_num)
                print(f"   📅 Найден месяц в тексте: {month_name} ({month_num})")
    
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Ф.И.О.' in row_str:
            header_row_idx = i
            break
    
    if header_row_idx is None:
        print(f"   ⚠️ Не найдена строка 'Ф.И.О.'")
        return {}, dept_name, months_found if months_found else None, years_found if years_found else None
    
    header_row = df.iloc[header_row_idx].values
    day_cols = []
    
    for i, val in enumerate(header_row):
        if pd.notna(val):
            day_num = extract_day_from_date(val)
            month_num = extract_month_from_date(val)
            year_num = extract_year_from_date(val)
            if month_num:
                months_found.add(month_num)
            if year_num:
                years_found.add(year_num)
            if day_num and 1 <= day_num <= 31:
                day_cols.append((i, day_num))
    
    print(f"   📅 Найдено {len(day_cols)} дней, месяц: {months_found}, год: {years_found}")
    
    result = {day: {} for _, day in day_cols}
    doctors_found = 0
    
    for i in range(header_row_idx + 1, len(df)):
        row = df.iloc[i].values
        if len(row) == 0 or pd.isna(row[0]) or str(row[0]).strip() == '':
            continue
        doctor = str(row[0]).strip()
        
        if not is_valid_doctor_name(doctor):
            continue
        
        doctors_found += 1
        for col_idx, day_num in day_cols:
            if col_idx < len(row) and pd.notna(row[col_idx]):
                val = str(row[col_idx]).strip()
                if val:
                    match = re.search(r'([АР])', val)
                    if match:
                        result[day_num][(doctor, dept_name)] = match.group(1)
    
    print(f"   👨‍⚕️ Найдено {doctors_found} врачей")
    
    return result, dept_name, months_found if months_found else None, years_found if years_found else None

def _parse_surgery_text(text, dept_name):
    """Парсит текст из PDF (корректная привязка дат)"""
    result = {}
    months_found = set()
    years_found = set()
    lines = text.split('\n')
    
    date_pattern = re.compile(r'(\d{2})/(\d{2})/(\d{4})')
    time_pattern = re.compile(r'(\d{2}):(\d{2})\s*[-—]\s*(\d{2}):(\d{2})')
    
    # Собираем все строки с данными
    parsed_entries = []
    current_day = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if 'Утверждаю' in line or 'Хирургическое' in line or 'График' in line:
            continue
        
        date_match = date_pattern.search(line)
        if date_match:
            day = int(date_match.group(1))
            month = int(date_match.group(2))
            year = int(date_match.group(3))
            months_found.add(month)
            years_found.add(year)
            current_day = day
            line = date_pattern.sub('', line).strip()
        
        if current_day is None:
            continue
        
        if not line:
            continue
        
        time_match = time_pattern.search(line)
        if time_match:
            time_str = f"{time_match.group(1)}/{time_match.group(3)}"
            normalized = normalize_time(time_str)
            doctor_part = line[:time_match.start()].strip()
            doctor = ' '.join(doctor_part.split())
            if doctor and is_valid_doctor_name(doctor):
                parsed_entries.append((current_day, doctor, normalized))
        else:
            alt_time_match = re.search(r'(\d{2}:\d{2})\s*[-—]\s*(\d{2}:\d{2})', line)
            if alt_time_match:
                time_str = f"{alt_time_match.group(1)[:2]}/{alt_time_match.group(2)[:2]}"
                normalized = normalize_time(time_str)
                doctor_part = line[:alt_time_match.start()].strip()
                doctor = ' '.join(doctor_part.split())
                if doctor and is_valid_doctor_name(doctor):
                    parsed_entries.append((current_day, doctor, normalized))
    
    # Корректная привязка дат
    day_doctors = {}
    for day, doctor, time_str in parsed_entries:
        if day not in day_doctors:
            day_doctors[day] = []
        day_doctors[day].append((doctor, time_str))
    
    sorted_days = sorted(day_doctors.keys())
    
    for i, day in enumerate(sorted_days):
        doctors = day_doctors[day]
        first_doctor = doctors[0]
        if day not in result:
            result[day] = {}
        result[day][(first_doctor[0], dept_name)] = first_doctor[1]
        
        if len(doctors) > 1 and i + 1 < len(sorted_days):
            next_day = sorted_days[i + 1]
            for doctor, time_str in doctors[1:]:
                if next_day not in result:
                    result[next_day] = {}
                result[next_day][(doctor, dept_name)] = time_str
        elif len(doctors) > 1 and i + 1 == len(sorted_days):
            for doctor, time_str in doctors[1:]:
                if day not in result:
                    result[day] = {}
                result[day][(doctor, dept_name)] = time_str
    
    return result, months_found, years_found

def parse_surgery_pdf(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name} (PDF)...")
    
    if not PDF_AVAILABLE:
        print(f"   ⚠️ PDF-библиотека не установлена")
        return {}, dept_name, None, None
    
    result = {}
    months_found = set()
    years_found = set()
    
    dept_name = "Хирургия"
    
    try:
        if PDF_LIB == 'pdfplumber':
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        print(f"   📄 Страница {page_num+1}: {len(text)} символов")
                        res, months, years = _parse_surgery_text(text, dept_name)
                        result.update(res)
                        months_found.update(months)
                        years_found.update(years)
        elif PDF_LIB == 'pypdf':
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    print(f"   📄 Страница {page_num+1}: {len(text)} символов")
                    res, months, years = _parse_surgery_text(text, dept_name)
                    result.update(res)
                    months_found.update(months)
                    years_found.update(years)
    except Exception as e:
        print(f"   ⚠️ Ошибка чтения PDF: {e}")
        return {}, dept_name, None, None
    
    print(f"   📅 Найдено {len(result)} дней с дежурствами, месяц: {months_found}, год: {years_found}")
    return result, dept_name, months_found, years_found

def parse_surgery_txt(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name} (TXT)...")
    
    encodings = ['utf-8-sig', 'cp1251', 'cp866', 'koi8-r', 'latin-1']
    lines = None
    
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                lines = f.readlines()
                break
        except UnicodeDecodeError:
            continue
    
    if lines is None:
        with open(file_path, 'rb') as f:
            raw = f.read()
            lines = raw.decode('utf-8', errors='ignore').splitlines()
    
    text = '\n'.join(lines)
    result, months, years = _parse_surgery_text(text, dept_name)
    
    print(f"   📅 Найдено {len(result)} дней с дежурствами, месяц: {months}, год: {years}")
    return result, dept_name, months, years

# ============================================================
# СОХРАНЕНИЕ В WORD
# ============================================================
def save_to_word(all_data, doctors_by_dept, sorted_depts, output_file, 
                 rukovoditel_text, ploshadka_text, month_num, year, month_name):
    """Сохраняет сводный график в Word с тремя колонками и шапкой"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime
    import calendar
    
    # Создаём документ
    doc = Document()
    
    # Настройка полей (узкие)
    section = doc.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    
    # ============================================================
    # ШАПКА (выравнивание по правому краю)
    # ============================================================
    ruk_parts = rukovoditel_text.split('|')
    ruk_dolzhnost = ruk_parts[0].strip() if len(ruk_parts) > 0 else ''
    ruk_fio = ruk_parts[1].strip() if len(ruk_parts) > 1 else ''
    
    # УТВЕРЖДАЮ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("УТВЕРЖДАЮ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Должность руководителя
    if ruk_dolzhnost:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(ruk_dolzhnost)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    
    # ФИО руководителя (с подчёркиванием)
    if ruk_fio:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"____________ {ruk_fio}")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    
    # Дата (сегодняшняя, в родительном падеже)
    today = datetime.datetime.now()
    month_names_ru = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    month_ru = month_names_ru.get(today.month, 'августа')
    date_str = f"{today.day} {month_ru} {today.year} г."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(date_str)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    
    # Пустая строка после шапки (одна)
    doc.add_paragraph()
    
    # ============================================================
    # ЗАГОЛОВОК (выравнивание по центру)
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("ГРАФИК ДЕЖУРСТВА ВРАЧЕЙ СТАЦИОНАРА ГОБУЗ МОКМЦ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(ploshadka_text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    month_names_upper = {
        1: 'ЯНВАРЬ', 2: 'ФЕВРАЛЬ', 3: 'МАРТ', 4: 'АПРЕЛЬ',
        5: 'МАЙ', 6: 'ИЮНЬ', 7: 'ИЮЛЬ', 8: 'АВГУСТ',
        9: 'СЕНТЯБРЬ', 10: 'ОКТЯБРЬ', 11: 'НОЯБРЬ', 12: 'ДЕКАБРЬ'
    }
    month_upper = month_names_upper.get(month_num, 'АВГУСТ')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"НА {month_upper} {year} Г.")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Пустая строка после заголовка (одна)
    doc.add_paragraph()
    
    # ============================================================
    # ДНИ НЕДЕЛИ
    # ============================================================
    _, last_day = calendar.monthrange(year, month_num)
    weekday_names = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
    weekdays = []
    for day in range(1, last_day + 1):
        date_obj = datetime.date(year, month_num, day)
        weekdays.append(weekday_names[date_obj.weekday()])
    
    # ============================================================
    # СБОРКА БЛОКОВ
    # ============================================================
    all_blocks = []
    for day in sorted(all_data.keys()):
        if 1 <= day <= last_day:
            day_name = weekdays[day-1]
            block_lines = [f"{day:02d}.{month_num:02d} {day_name}"]
            
            for dept in sorted_depts:
                dept_doctors = sorted(doctors_by_dept.get(dept, []), key=lambda x: x[0])
                abbr = get_dept_abbr(dept)
                
                for doctor, dept_key in dept_doctors:
                    if (doctor, dept_key) in all_data[day]:
                        value = all_data[day][(doctor, dept_key)]
                        block_lines.append(f"{doctor} ({abbr}) {value}")
            
            all_blocks.append(block_lines)
    
    # ============================================================
    # РАСПРЕДЕЛЕНИЕ ПО СТРАНИЦАМ
    # ============================================================
    def calc_rows_for_col(blocks):
        total = 0
        for block_idx, block in enumerate(blocks):
            is_last = (block_idx == len(blocks) - 1)
            total += len(block)
            if not is_last:
                total += 1
        return total
    
    FIRST_PAGE_MAX_ROWS = 48
    OTHER_PAGE_MAX_ROWS = 59
    
    def distribute_blocks_with_limit(blocks, first_max, other_max):
        pages = []
        current_page = [[], [], []]
        col_heights = [0, 0, 0]
        current_col = 0
        is_first_page = True
        
        for block in blocks:
            block_rows = len(block) + 1
            
            if is_first_page:
                max_rows = first_max
            else:
                max_rows = other_max
            
            if col_heights[current_col] + block_rows <= max_rows:
                current_page[current_col].append(block)
                col_heights[current_col] += block_rows
            else:
                col_found = False
                for next_col in range(current_col + 1, 3):
                    if col_heights[next_col] + block_rows <= max_rows:
                        current_col = next_col
                        current_page[current_col].append(block)
                        col_heights[current_col] += block_rows
                        col_found = True
                        break
                
                if not col_found:
                    if any(len(col) > 0 for col in current_page):
                        pages.append(current_page)
                    
                    current_page = [[], [], []]
                    col_heights = [0, 0, 0]
                    current_col = 0
                    is_first_page = False
                    
                    current_page[current_col].append(block)
                    col_heights[current_col] += block_rows
        
        if any(len(col) > 0 for col in current_page):
            pages.append(current_page)
        
        return pages
    
    pages = distribute_blocks_with_limit(all_blocks, FIRST_PAGE_MAX_ROWS, OTHER_PAGE_MAX_ROWS)
    
    if not pages:
        pages = [[[], [], []]]
    
    # ============================================================
    # ЗАПОЛНЕНИЕ КОЛОНОК
    # ============================================================
    def fill_column(table, col_idx, blocks, start_row, max_rows):
        row = start_row
        for block_idx, block in enumerate(blocks):
            if row >= max_rows:
                break
                
            cell = table.cell(row, col_idx)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            run = p.add_run(block[0])
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            row += 1
            
            for line in block[1:]:
                if row >= max_rows:
                    break
                cell = table.cell(row, col_idx)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                run = p.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                row += 1
            
            is_last_block = (block_idx == len(blocks) - 1)
            if not is_last_block and row < max_rows:
                cell = table.cell(row, col_idx)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.add_run('')
                row += 1
        
        return row
    
    # ============================================================
    # СОЗДАНИЕ ТАБЛИЦ
    # ============================================================
    first_page = True
    for page_idx, page in enumerate(pages):
        if not any(len(col) > 0 for col in page):
            continue
        
        if not first_page:
            doc.add_page_break()
        
        first_page = False
        
        max_rows = max(
            calc_rows_for_col(page[0]),
            calc_rows_for_col(page[1]),
            calc_rows_for_col(page[2])
        )
        table_height = max(max_rows, 3)
        
        table = doc.add_table(rows=table_height, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        fill_column(table, 0, page[0], 0, table_height)
        fill_column(table, 1, page[1], 0, table_height)
        fill_column(table, 2, page[2], 0, table_height)
        
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(2.5)
    
    doc.save(output_file)
    print(f"💾 Word сохранён: {output_file}")

# ============================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================
# ============================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================
def build_master_schedule(input_folder, rukovoditel_text, ploshadka_text, 
                          selected_month, selected_year):
    """Собирает сводный график из файлов в папке"""
    print(f"\n📊 Начинаем сборку сводного графика...")
    print(f"📁 Папка: {input_folder}")
    print(f"📅 Месяц: {selected_month}, Год: {selected_year}")
    print("=" * 60)
    
    folder = Path(input_folder)
    
    output_folder = folder / "Сводный график"
    output_folder.mkdir(exist_ok=True)
    
    all_files = []
    for ext in ['*.xlsx', '*.xls', '*.csv', '*.pdf', '*.txt']:
        all_files.extend(folder.glob(ext))
    
    all_files = [f for f in all_files if not f.name.startswith('~')]
    all_files = [f for f in all_files if 'сводный' not in f.name.lower()]
    
    if not all_files:
        print("❌ Не найдено ни одного файла с графиками!")
        return False, [], "❌ Не найдено ни одного файла с графиками!"
    
    print(f"\n📄 Найдено {len(all_files)} файлов:")
    for f in all_files:
        print(f"   - {f.name}")
    
    print("\n" + "=" * 60)
    
    all_data = {}
    doctors_by_dept = {}
    dept_order_map = {}
    files_with_errors = []
    
    for file_path in all_files:
        print(f"\n{'='*60}")
        print(f"📂 Обрабатываем: {file_path.name}")
        
        df_temp = read_excel_or_csv(file_path)
        if df_temp is not None:
            dept_name = detect_department(df_temp, file_path)
        else:
            dept_name = Path(file_path).stem.replace('_', ' ')
        
        print(f"   🏥 Отделение: {dept_name}")
        print(f"   📛 Сокращение: {get_dept_abbr(dept_name)}")
        
        if dept_name == 'ОАР' or 'ОАР' in file_path.name:
            data, detected_dept, months, years = parse_oar_graph(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        elif dept_name == 'Ответственные по стационару' or 'ответственн' in dept_name.lower() or 'Ответственные' in file_path.name:
            data, detected_dept, months, years = parse_responsible_graph(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        elif file_path.suffix.lower() == '.pdf':
            data, detected_dept, months, years = parse_surgery_pdf(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        elif file_path.suffix.lower() == '.txt':
            data, detected_dept, months, years = parse_surgery_txt(file_path, dept_name)
            if not data:
                print(f"   ⚠️ Не удалось прочитать TXT, пропускаем")
                continue
        else:
            data, detected_dept, months, years = parse_standard_graph(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        
        if not data:
            print(f"   ⚠️ Данные не найдены для {dept_name}")
            continue
        
        if months:
            if selected_month not in months:
                print(f"   ⚠️ Месяц в файле ({months}) не соответствует выбранному ({selected_month})")
                files_with_errors.append((file_path.name, f"Месяц: {months}, ожидался: {selected_month}"))
                continue
        
        if years:
            if selected_year not in years:
                print(f"   ⚠️ Год в файле ({years}) не соответствует выбранному ({selected_year})")
                files_with_errors.append((file_path.name, f"Год: {years}, ожидался: {selected_year}"))
                continue
        
        for day, doctors in data.items():
            if day not in all_data:
                all_data[day] = {}
            all_data[day].update(doctors)
            
            if dept_name not in doctors_by_dept:
                doctors_by_dept[dept_name] = []
            for doctor_key in doctors.keys():
                if doctor_key not in doctors_by_dept[dept_name]:
                    doctors_by_dept[dept_name].append(doctor_key)
        
        dept_order_map[dept_name] = get_dept_order(dept_name)
        
        print(f"   ✅ Обработано {len(data)} дней, {len(doctors_by_dept.get(dept_name, []))} врачей")
    
    if not all_data:
        print("❌ Данные не найдены ни в одном файле!")
        month_name_ru = MONTHS_RU_NOMINATIVE.get(selected_month, '')
        return False, files_with_errors, f"❌ Данные за {month_name_ru} {selected_year} не найдены! Проверьте месяц в файлах"
    
    if files_with_errors:
        print("\n" + "=" * 60)
        print("⚠️⚠️⚠️ Файлы, исключённые из сводного графика (не соответствуют выбранному месяцу/году):")
        print("-" * 60)
        for fname, error in files_with_errors:
            print(f"   ❌ {fname}")
            print(f"      {error}")
        print("-" * 60)
        print(f"   Всего исключено: {len(files_with_errors)} файлов")
        print("=" * 60)
    
    sorted_depts = sorted(
        doctors_by_dept.keys(),
        key=lambda d: (dept_order_map.get(d, 999), d)
    )
    
    print("\n" + "=" * 60)
    print(f"✅ Всего найдено {len(all_data)} дней с дежурствами")
    print(f"📊 Найдено отделений: {len(doctors_by_dept)}")
    for dept in sorted_depts:
        abbr = get_dept_abbr(dept)
        print(f"   {dept} ({abbr}): {len(doctors_by_dept[dept])} врачей")
    
    month_name = MONTHS_RU_UPPER.get(selected_month, 'АВГУСТ')
    
    _, last_day = calendar.monthrange(selected_year, selected_month)
    weekday_names = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
    weekdays = []
    for day in range(1, last_day + 1):
        date_obj = datetime.date(selected_year, selected_month, day)
        weekdays.append(weekday_names[date_obj.weekday()])
    
    output_lines = []
    for day in sorted(all_data.keys()):
        if 1 <= day <= last_day:
            day_name = weekdays[day-1]
            output_lines.append(f"{day:02d}.{selected_month:02d} {day_name}")
            
            for dept in sorted_depts:
                dept_doctors = sorted(doctors_by_dept.get(dept, []), key=lambda x: x[0])
                abbr = get_dept_abbr(dept)
                
                for doctor, dept_key in dept_doctors:
                    if (doctor, dept_key) in all_data[day]:
                        value = all_data[day][(doctor, dept_key)]
                        output_lines.append(f"{doctor} ({abbr}) {value}")
            
            output_lines.append("")
    
    month_lower = MONTHS_RU.get(selected_month, 'августа')
    txt_file = output_folder / f"сводный_график_{month_lower}_{selected_year}.txt"
    with open(txt_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))
    
    print(f"\n💾 TXT сохранён: {txt_file}")
    
    docx_file = output_folder / f"сводный_график_{month_lower}_{selected_year}.docx"
    save_to_word(all_data, doctors_by_dept, sorted_depts, docx_file,
                 rukovoditel_text, ploshadka_text, selected_month, selected_year, month_name)
    
    print(f"\n✅ Готово!")
    print(f"   📄 TXT: {txt_file}")
    print(f"   📝 Word: {docx_file}")
    
    return True, files_with_errors, ""  # ← ВОЗВРАЩАЕМ 3 ЗНАЧЕНИЯ
# ============================================================
# ГЛАВНОЕ ПРИЛОЖЕНИЕ (CustomTkinter)
# ============================================================
import ctypes

class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        # Устанавливаем иконку для окна
        try:
            icon_path = os.path.join(get_base_path(), "icon.ico")
            if os.path.exists(icon_path):
                # Для окна
                self.iconbitmap(icon_path)
                # Для панели задач (Windows)
                ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("myappid")
        except:
            pass
        
        # ... остальной код ...
        
        self.title("Сводный график дежурств")
        self.geometry("700x750")
        self.minsize(650, 700)
        
        # Переменные
        self.folder_path = ctk.StringVar(value="")
        self.rukovoditel = ctk.StringVar()
        self.ploshadka = ctk.StringVar()
        self.month_var = ctk.StringVar()
        self.year_var = ctk.StringVar()
        
        # Загружаем справочники
        self.rukovoditeli_list = load_spravochnik('spravochnik_rukovoditeli.txt')
        self.ploshadki_list = load_spravochnik('spravochnik_ploshadki.txt')
        
        # Определяем год по умолчанию
        now = datetime.datetime.now()
        default_year = now.year
        if now.month == 12 and now.day > 15:
            default_year += 1
        self.year_var.set(str(default_year))
        
        # Создаём интерфейс
        self.create_widgets()
    
    def create_widgets(self):
        # ============================================================
        # ЗАГОЛОВОК
        # ============================================================
        self.label_title = ctk.CTkLabel(
            self, 
            text="📊 Сводный график дежурств",
            font=ctk.CTkFont(size=22, weight="bold")
        )
        self.label_title.pack(pady=(20, 15))
        
        # ============================================================
        # ПАПКА С ГРАФИКАМИ
        # ============================================================
        self.frame_folder = ctk.CTkFrame(self)
        self.frame_folder.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(
            self.frame_folder, 
            text="📁 Папка с графиками:",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(anchor="w", padx=5, pady=(5, 0))
        
        folder_row = ctk.CTkFrame(self.frame_folder, fg_color="transparent")
        folder_row.pack(fill="x", padx=5, pady=5)
        
        self.entry_folder = ctk.CTkEntry(
            folder_row, 
            textvariable=self.folder_path,
            placeholder_text="Выберите папку с файлами графиков...",
            width=500
        )
        self.entry_folder.pack(side="left", fill="x", expand=True, padx=(0, 5))
        
        self.btn_folder = ctk.CTkButton(
            folder_row, 
            text="Обзор...", 
            command=self.select_folder,
            width=80
        )
        self.btn_folder.pack(side="right")
        
        # ============================================================
        # ПЕРИОД
        # ============================================================
        self.frame_date = ctk.CTkFrame(self)
        self.frame_date.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(
            self.frame_date, 
            text="📅 Период:",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(anchor="w", padx=5, pady=(5, 0))
        
        date_row = ctk.CTkFrame(self.frame_date, fg_color="transparent")
        date_row.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(date_row, text="Месяц:").pack(side="left", padx=(0, 10))
        
        months_list = [MONTHS_RU_NOMINATIVE[i] for i in range(1, 13)]
        self.combo_month = ctk.CTkComboBox(
            date_row, 
            values=months_list,
            width=150
        )
        self.combo_month.pack(side="left", padx=(0, 20))
        self.combo_month.set("Август")
        
        ctk.CTkLabel(date_row, text="Год:").pack(side="left", padx=(0, 10))
        
        years_list = [str(y) for y in range(2024, 2031)]
        self.combo_year = ctk.CTkComboBox(
            date_row, 
            values=years_list,
            width=100
        )
        self.combo_year.pack(side="left")
        self.combo_year.set(str(datetime.datetime.now().year))
        
        # ============================================================
        # РУКОВОДИТЕЛЬ
        # ============================================================
        self.frame_ruk = ctk.CTkFrame(self)
        self.frame_ruk.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(
            self.frame_ruk, 
            text="👤 Руководитель:",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(anchor="w", padx=5, pady=(5, 0))
        
        if self.rukovoditeli_list:
            self.combo_ruk = ctk.CTkComboBox(
                self.frame_ruk,
                values=self.rukovoditeli_list,
                width=500
            )
            self.combo_ruk.pack(padx=5, pady=5)
            if len(self.rukovoditeli_list) > 1:
                self.combo_ruk.set(self.rukovoditeli_list[1])
            else:
                self.combo_ruk.set(self.rukovoditeli_list[0])
        else:
            ctk.CTkLabel(
                self.frame_ruk, 
                text="⚠️ Нет данных в справочнике руководителей"
            ).pack(padx=5, pady=5)
        
        # ============================================================
        # ПЛОЩАДКА
        # ============================================================
        self.frame_plo = ctk.CTkFrame(self)
        self.frame_plo.pack(fill="x", padx=20, pady=5)
        
        ctk.CTkLabel(
            self.frame_plo, 
            text="🏥 Площадка:",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(anchor="w", padx=5, pady=(5, 0))
        
        if self.ploshadki_list:
            self.combo_plo = ctk.CTkComboBox(
                self.frame_plo,
                values=self.ploshadki_list,
                width=500
            )
            self.combo_plo.pack(padx=5, pady=5)
            if self.ploshadki_list:
                self.combo_plo.set(self.ploshadki_list[0])
        else:
            ctk.CTkLabel(
                self.frame_plo, 
                text="⚠️ Нет данных в справочнике площадок"
            ).pack(padx=5, pady=5)
        
        # ============================================================
        # СПИСОК ФАЙЛОВ
        # ============================================================
        self.frame_files = ctk.CTkFrame(self)
        self.frame_files.pack(fill="both", expand=True, padx=20, pady=5)
        
        ctk.CTkLabel(
            self.frame_files, 
            text="📄 Файлы в папке:",
            font=ctk.CTkFont(size=13, weight="bold")
        ).pack(anchor="w", padx=5, pady=(5, 0))
        
        self.files_listbox = ctk.CTkTextbox(
            self.frame_files,
            height=80,
            font=ctk.CTkFont(size=11)
        )
        self.files_listbox.pack(fill="both", expand=True, padx=5, pady=5)
        
        btn_refresh = ctk.CTkButton(
            self.frame_files,
            text="🔄 Обновить список",
            command=self.refresh_files_list,
            width=150
        )
        btn_refresh.pack(pady=5)
        
        # ============================================================
        # КНОПКА ЗАПУСКА
        # ============================================================
        self.btn_run = ctk.CTkButton(
            self,
            text="🚀 Составить сводный график",
            font=ctk.CTkFont(size=14, weight="bold"),
            fg_color="#4CAF50",
            hover_color="#388E3C",
            height=45,
            command=self.run
        )
        self.btn_run.pack(pady=20, padx=20, fill="x")
        
        # ============================================================
        # СТАТУС (нижняя часть) — УВЕЛИЧЕННЫЙ
        # ============================================================
        self.status_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.status_frame.pack(fill="x", padx=20, pady=(0, 15))

        self.status_textbox = ctk.CTkTextbox(
            self.status_frame,
            height=80,
            font=ctk.CTkFont(size=13),
            wrap="word"
        )
        self.status_textbox.pack(fill="x", pady=5)
        self.status_textbox.insert("0.0", "✅ Готов к работе")
        self.status_textbox.configure(state="disabled")

        btn_clear_status = ctk.CTkButton(
            self.status_frame,
            text="🗑️ Очистить",
            command=self.clear_status,
            width=80,
            height=28,
            fg_color="#607D8B",
            hover_color="#455A64"
        )
        btn_clear_status.pack(pady=(0, 5))

    def set_status(self, message, status_type="info", details=None):
        """
        Обновляет статус в нижней части окна
        status_type: "info" (синий), "success" (зелёный), "error" (красный), "warning" (оранжевый)
        details: дополнительная информация (список файлов и т.д.)
        """
        colors = {
            "info": "#4FC3F7",
            "success": "#66BB6A",
            "error": "#EF5350",
            "warning": "#FFA726"
        }
        color = colors.get(status_type, "#FFFFFF")
        
        # Разрешаем редактирование
        self.status_textbox.configure(state="normal")
        self.status_textbox.delete("0.0", "end")
        
        # Вставляем основное сообщение с цветом
        self.status_textbox.insert("0.0", message + "\n", status_type)
        self.status_textbox.tag_config(status_type, foreground=color)
        
        # Если есть детали (список файлов) — добавляем
        if details:
            self.status_textbox.insert("end", "\n" + details, "details")
            self.status_textbox.tag_config("details", foreground="#B0BEC5")  # ← убрали font
        
        self.status_textbox.configure(state="disabled")
        self.update()
    
    def clear_status(self):
        """Очищает статус"""
        self.status_textbox.configure(state="normal")
        self.status_textbox.delete("0.0", "end")
        self.status_textbox.insert("0.0", "✅ Готов к работе")
        self.status_textbox.configure(state="disabled")
        self.update()

    def show_excluded_files(self, files_with_errors):
        """Показывает список исключённых файлов в статусе"""
        if not files_with_errors:
            return
        
        count = len(files_with_errors)
        
        # Показываем все файлы, если их не больше 10
        if count <= 10:
            files_list = "\n".join([f"  • {fname} — {error}" for fname, error in files_with_errors])
        else:
            files_list = "\n".join([f"  • {fname} — {error}" for fname, error in files_with_errors[:10]])
            files_list += f"\n  ... и ещё {count - 10} файлов"
        
        details = f"Исключённые файлы ({count}):\n{files_list}"
        self.set_status(
            f"⚠️ {count} файлов исключены (не соответствуют месяцу)",
            "warning",
            details
        )

    def refresh_files_list(self):
        """Обновляет список файлов в папке"""
        folder = self.folder_path.get().strip()
        self.files_listbox.delete("0.0", "end")
        
        if not folder:
            self.files_listbox.insert("0.0", "⚠️ Папка не выбрана")
            self.set_status("⚠️ Папка не выбрана", "warning")  # ← вместо status_label
            return
        
        if not os.path.exists(folder):
            self.files_listbox.insert("0.0", "⚠️ Папка не существует")
            self.set_status("⚠️ Папка не существует", "warning")  # ← вместо status_label
            return
        
        try:
            files = os.listdir(folder)
            graph_files = sorted([f for f in files if f.endswith(('.xlsx', '.xls', '.csv', '.pdf', '.txt'))])
            if graph_files:
                for f in graph_files:
                    self.files_listbox.insert("end", f + "\n")
                self.set_status(f"✅ Найдено {len(graph_files)} файлов", "success")  # ← вместо status_label
            else:
                self.files_listbox.insert("0.0", "⚠️ Нет файлов графиков")
                self.set_status("⚠️ В папке нет файлов графиков", "warning")  # ← вместо status_label
        except Exception as e:
            self.files_listbox.insert("0.0", f"❌ Ошибка: {str(e)}")
            self.set_status(f"❌ Ошибка: {str(e)}", "error")  # ← вместо status_label
    
    def select_folder(self):
        folder = filedialog.askdirectory(title="Выберите папку с графиками")
        if folder:
            self.folder_path.set(folder)
            self.refresh_files_list()
    
    def run(self):
        folder = self.folder_path.get().strip()
        if not folder:
            self.set_status("❌ Ошибка: выберите папку с графиками!", "error")
            return
        
        if not os.path.exists(folder):
            self.set_status(f"❌ Папка не существует: {folder}", "error")
            return
        
        ruk_text = self.combo_ruk.get().strip()
        if not ruk_text:
            self.set_status("❌ Ошибка: выберите руководителя!", "error")
            return
        
        plo_text = self.combo_plo.get().strip()
        if not plo_text:
            self.set_status("❌ Ошибка: выберите площадку!", "error")
            return
        
        # Парсим месяц
        month_str = self.combo_month.get().strip()
        month_names = {v: k for k, v in MONTHS_RU_NOMINATIVE.items()}
        if month_str in month_names:
            month_num = month_names[month_str]
        else:
            month_str_lower = month_str.lower()
            month_names_ru = {v.lower(): k for k, v in MONTHS_RU.items()}
            if month_str_lower in month_names_ru:
                month_num = month_names_ru[month_str_lower]
            else:
                self.set_status(f"❌ Неверный месяц: {month_str}", "error")
                return
        
        # Парсим год
        try:
            year_num = int(self.combo_year.get().strip())
        except:
            self.set_status("❌ Неверный формат года!", "error")
            return
        
        self.set_status("⏳ Обработка... Пожалуйста, подождите", "info")
        self.btn_run.configure(state="disabled", text="⏳ Обработка...")
        self.update()
        
        try:
            # ВЫЗЫВАЕМ ФУНКЦИЮ И ПОЛУЧАЕМ ТРИ ЗНАЧЕНИЯ
            success, files_with_errors, error_message = build_master_schedule(
                folder, ruk_text, plo_text, month_num, year_num
            )
            
            if success:
                self.set_status("✅ Сводный график успешно создан! Файлы в папке 'Сводный график'", "success")
                if files_with_errors:
                    self.show_excluded_files(files_with_errors)
            else:
                # Если есть явное сообщение об ошибке — показываем его
                if error_message:
                    self.set_status(error_message, "error")
                elif files_with_errors:
                    self.show_excluded_files(files_with_errors)
                else:
                    self.set_status("❌ Ошибка при создании графика", "error")
            
        except Exception as e:
            self.set_status(f"❌ Критическая ошибка: {str(e)}", "error")
            import traceback
            print(traceback.format_exc())
        
        finally:
            self.btn_run.configure(state="normal", text="🚀 Составить сводный график")
            self.update()

# ============================================================
# ЗАПУСК
# ============================================================
if __name__ == '__main__':
    app = App()
    app.mainloop()