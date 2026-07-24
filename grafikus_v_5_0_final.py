import pandas as pd
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import datetime
import os
import calendar
import sys



def get_base_path():
    """Возвращает путь к папке, где находится исполняемый файл (или скрипт)"""
    if getattr(sys, 'frozen', False):
        # Запущено как .exe
        return os.path.dirname(sys.executable)
    else:
        # Запущено как .py
        return os.path.dirname(os.path.abspath(__file__))

# ============================================================
# ПЫТАЕМСЯ ИМПОРТИРОВАТЬ БИБЛИОТЕКУ ДЛЯ PDF
# ============================================================
PDF_AVAILABLE = False
PDF_LIB = None

try:
    import pdfplumber
    PDF_AVAILABLE = True
    PDF_LIB = 'pdfplumber'
except ImportError:
    try:
        from pypdf import PdfReader
        PDF_AVAILABLE = True
        PDF_LIB = 'pypdf'
    except ImportError:
        print("⚠️ Ни одна PDF-библиотека не установлена.")

# ============================================================
# СПРАВОЧНИКИ
# ============================================================
def load_spravochnik(filename):
    """Загружает справочник из текстового файла (рядом с .exe)"""
    base_path = get_base_path()
    file_path = os.path.join(base_path, filename)
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]
        return lines
    except FileNotFoundError:
        # Если файла нет, создаём с данными по умолчанию
        print(f"⚠️ Файл {filename} не найден, создаём с данными по умолчанию")
        return []

def save_spravochnik(filename, data):
    """Сохраняет справочник в текстовый файл (рядом с .exe)"""
    base_path = get_base_path()
    file_path = os.path.join(base_path, filename)
    
    with open(file_path, 'w', encoding='utf-8') as f:
        for item in data:
            f.write(item + '\n')

# Загружаем справочники
RUKOVODITELI = load_spravochnik('spravochnik_rukovoditeli.txt')
PLOSHADKI = load_spravochnik('spravochnik_ploshadki.txt')

# Если справочники пустые, создаём с данными по умолчанию
if not RUKOVODITELI:
    RUKOVODITELI = [
        'Главный врач ГОБУЗ МОКМЦ|Тарбаев Е.Ю.',
        'Зам. гл. врача по медицинской части ГОБУЗ МОКМЦ|Гредягин С.С.',
        'Зам. гл. врача по терапии ГОБУЗ МОКМЦ|Колосова О.Л.'
    ]
    save_spravochnik('spravochnik_rukovoditeli.txt', RUKOVODITELI)

if not PLOSHADKI:
    PLOSHADKI = ['Ломоносова, д.18', 'Володарского, д.18', 'Перинатальный центр', 'Родильный дом']
    save_spravochnik('spravochnik_ploshadki.txt', PLOSHADKI)

# Месяцы на русском
MONTHS_RU = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
}

# Добавьте после MONTHS_RU:
MONTHS_RU_NOMINATIVE = {
    1: 'Январь', 2: 'Февраль', 3: 'Март', 4: 'Апрель',
    5: 'Май', 6: 'Июнь', 7: 'Июль', 8: 'Август',
    9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'
}

MONTHS_RU_UPPER = {
    1: 'ЯНВАРЬ', 2: 'ФЕВРАЛЬ', 3: 'МАРТ', 4: 'АПРЕЛЬ',
    5: 'МАЙ', 6: 'ИЮНЬ', 7: 'ИЮЛЬ', 8: 'АВГУСТ',
    9: 'СЕНТЯБРЬ', 10: 'ОКТЯБРЬ', 11: 'НОЯБРЬ', 12: 'ДЕКАБРЬ'
}

# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================
def normalize_time(time_str):
    if not time_str or time_str == '':
        return time_str
    
    time_str = time_str.strip()
    
    match = re.match(r'(\d{1,2}):(\d{2})\s*[-—]\s*(\d{1,2}):(\d{2})', time_str)
    if match:
        h1, m1, h2, m2 = match.groups()
        return f"{int(h1):02d}:{m1}-{int(h2):02d}:{m2}"
    
    match = re.match(r'(\d{1,2})\s*[/]\s*(\d{1,2})', time_str)
    if match:
        h1, h2 = match.groups()
        return f"{int(h1):02d}-{int(h2):02d}"
    
    match = re.match(r'(\d{1,2})\s*[-—]\s*(\d{1,2})', time_str)
    if match:
        h1, h2 = match.groups()
        return f"{int(h1):02d}-{int(h2):02d}"
    
    match = re.match(r'0-(\d{1,2})', time_str)
    if match:
        h2 = match.group(1)
        return f"00-{int(h2):02d}"
    
    return time_str

def read_excel_or_csv(file_path):
    file_path = Path(file_path)
    
    if not file_path.exists():
        return None
    
    if file_path.suffix.lower() in ['.xlsx', '.xls']:
        try:
            df = pd.read_excel(file_path, header=None)
            return df
        except Exception as e:
            print(f"   ⚠️ Ошибка чтения Excel: {e}")
            return None
    
    elif file_path.suffix.lower() == '.csv':
        try:
            df = pd.read_csv(file_path, sep=';', encoding='utf-8-sig', header=None)
            return df
        except:
            try:
                df = pd.read_csv(file_path, sep=';', encoding='cp1251', header=None)
                return df
            except:
                try:
                    df = pd.read_csv(file_path, sep=';', encoding='latin-1', header=None)
                    return df
                except:
                    return None
    
    return None

def get_dept_abbr(dept_name):
    DEPT_ABBR = {
        'Ответственные по стационару': 'ОТВ',
        'Приемное отделение': 'ПО',
        'ОАР': 'ОАР',
        'Хирургия': 'ХИР',
        'Гинекология': 'ГИН',
        'Урология': 'УРО',
        'Травматология': 'ТРАВМ',
        'Травмпункт': 'ТР П',
    }
    
    if dept_name in DEPT_ABBR:
        return DEPT_ABBR[dept_name]
    
    for key, abbr in DEPT_ABBR.items():
        if key.lower() in dept_name.lower() or dept_name.lower() in key.lower():
            return abbr
    
    words = dept_name.split()
    if len(words) >= 2:
        return ''.join(w[0].upper() for w in words[:2])
    return dept_name[:4].upper()

def get_dept_order(dept_name):
    DEPARTMENT_ORDER = [
        'Ответственные по стационару',
        'Приемное отделение',
        'ОАР',
        'Хирургия',
        'Гинекология',
        'Урология',
        'Травматология',
        'Травмпункт',
    ]
    for i, dept in enumerate(DEPARTMENT_ORDER):
        if dept.lower() in dept_name.lower() or dept_name.lower() in dept.lower():
            return i
    return len(DEPARTMENT_ORDER)

def extract_day_from_date(val):
    val_str = str(val).strip()
    
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', val_str)
    if match:
        return int(match.group(3))
    
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', val_str)
    if match:
        return int(match.group(1))
    
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', val_str)
    if match:
        return int(match.group(1))
    
    match = re.search(r'\b([1-9]|[12][0-9]|3[01])\b', val_str)
    if match:
        return int(match.group(1))
    
    return None

def extract_month_from_date(val):
    val_str = str(val).strip()
    
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', val_str)
    if match:
        return int(match.group(2))
    
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', val_str)
    if match:
        return int(match.group(2))
    
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', val_str)
    if match:
        return int(match.group(2))
    
    return None

def extract_year_from_date(val):
    val_str = str(val).strip()
    
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', val_str)
    if match:
        return int(match.group(1))
    
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', val_str)
    if match:
        return int(match.group(3))
    
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', val_str)
    if match:
        return int(match.group(3))
    
    return None

def is_valid_doctor_name(name):
    if not name or not isinstance(name, str):
        return False
    
    name = name.strip()
    if not name:
        return False
    
    headers = ['Дата', 'Дежурный врач', 'Сб', 'Вс', 'Пн', 'Вт', 'Ср', 'Чт', 'Пт',
               'График', 'Утверждаю', 'Заведующий', 'Ф.И.О.', 'цех', 'отделение',
               '2026', '2025', 'Август', 'Март']
    for h in headers:
        if h in name:
            return False
    
    if re.search(r'\d{4}-\d{2}-\d{2}', name):
        return False
    
    if len(name) > 50:
        return False
    
    if not re.search(r'[а-яА-ЯёЁ]', name):
        return False
    
    return True

# ============================================================
# ПАРСЕРЫ
# ============================================================
def detect_department(df, file_path):
    file_stem = Path(file_path).stem.lower()
    
    header_text = ''
    for i in range(min(10, len(df))):
        row = df.iloc[i].values
        row_str = ' '.join(str(v) for v in row if pd.notna(v))
        header_text += row_str + ' '
    
    header_text = header_text.lower()
    
    DEPARTMENT_ORDER = [
        'Ответственные по стационару',
        'Приемное отделение',
        'ОАР',
        'Хирургия',
        'Гинекология',
        'Урология',
        'Травматология',
        'Травмпункт',
    ]
    
    for dept in DEPARTMENT_ORDER:
        dept_clean = re.sub(r'[_\s]+', ' ', dept.lower())
        clean_filename = re.sub(r'графики?_дежурных?_врачей?_', '', file_stem)
        clean_filename = re.sub(r'[_\s]+', ' ', clean_filename).strip()
        if dept_clean in clean_filename or clean_filename in dept_clean:
            return dept
    
    dept_patterns = {
        'Ответственные по стационару': ['ответственн', 'ответсвтвенн', 'стационар'],
        'Приемное отделение': ['приемн', 'приёмн', 'приемного'],
        'ОАР': ['оар', 'реаниматолог', 'анестезиолог', 'анестезия'],
        'Хирургия': ['хирург', 'хирургическ'],
        'Гинекология': ['гинеколог', 'гинекологическ'],
        'Урология': ['уролог', 'урологическ'],
        'Травматология': ['травматолог', 'травматологическ'],
        'Травмпункт': ['травмпункт', 'травм пункт'],
    }
    
    for dept, patterns in dept_patterns.items():
        for pattern in patterns:
            if pattern in header_text:
                return dept
    
    return Path(file_path).stem.replace('_', ' ')

def parse_standard_graph(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name}...")
    
    df = read_excel_or_csv(file_path)
    if df is None:
        print(f"   ⚠️ Не удалось прочитать файл")
        return {}, dept_name, None, None
    
    date_row_idx = None
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Дата' in row_str and re.search(r'\d{4}-\d{2}-\d{2}|\d{2}\.\d{2}\.\d{4}', row_str):
            date_row_idx = i
            break
    
    if date_row_idx is None:
        for i, row in df.iterrows():
            row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
            numbers = re.findall(r'\d+', row_str)
            if len(numbers) >= 10:
                date_row_idx = i
                break
    
    if date_row_idx is None:
        print(f"   ⚠️ Не найдена строка с датами")
        return {}, dept_name, None, None
    
    date_row = df.iloc[date_row_idx].values
    day_cols = []
    months_found = set()
    years_found = set()
    
    for i, val in enumerate(date_row):
        if pd.notna(val):
            day_num = extract_day_from_date(val)
            month_num = extract_month_from_date(val)
            year_num = extract_year_from_date(val)
            if month_num:
                months_found.add(month_num)
            if year_num:
                years_found.add(year_num)
            if day_num and 1 <= day_num <= 31:
                day_cols.append((i, day_num))
    
    print(f"   📅 Найдено {len(day_cols)} дней, месяц: {months_found}, год: {years_found}")
    
    header_row_idx = None
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Дежурный врач' in row_str or 'дежурный' in row_str.lower():
            header_row_idx = i
            break
    
    if header_row_idx is None:
        print(f"   ⚠️ Не найдена строка 'Дежурный врач'")
        return {}, dept_name, None, None
    
    result = {day: {} for _, day in day_cols}
    doctors_found = 0
    
    for i in range(header_row_idx + 1, len(df)):
        row = df.iloc[i].values
        if len(row) == 0 or pd.isna(row[0]) or str(row[0]).strip() == '':
            continue
        doctor = str(row[0]).strip()
        
        if not is_valid_doctor_name(doctor):
            continue
        
        if doctor == '' or doctor == 'Заведующий' or 'Заведующий' in doctor:
            break
        
        doctors_found += 1
        for col_idx, day_num in day_cols:
            if col_idx < len(row) and pd.notna(row[col_idx]):
                val = str(row[col_idx]).strip()
                if val and val not in ['', 'nan', 'None']:
                    if re.search(r'\d+[-/]\d+|\d+:\d+', val):
                        normalized = normalize_time(val)
                        result[day_num][(doctor, dept_name)] = normalized
    
    print(f"   👨‍⚕️ Найдено {doctors_found} врачей")
    
    return result, dept_name, months_found, years_found

def parse_responsible_graph(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name}...")
    
    df = read_excel_or_csv(file_path)
    if df is None:
        print(f"   ⚠️ Не удалось прочитать файл")
        return {}, dept_name, None, None
    
    date_row_idx = None
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Дата' in row_str and re.search(r'\d{4}-\d{2}-\d{2}|\d{2}\.\d{2}\.\d{4}', row_str):
            date_row_idx = i
            break
    
    if date_row_idx is None:
        for i, row in df.iterrows():
            row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
            numbers = re.findall(r'\d+', row_str)
            if len(numbers) >= 10:
                date_row_idx = i
                break
    
    if date_row_idx is None:
        print(f"   ⚠️ Не найдена строка с датами")
        return {}, dept_name, None, None
    
    date_row = df.iloc[date_row_idx].values
    day_cols = []
    months_found = set()
    years_found = set()
    
    for i, val in enumerate(date_row):
        if pd.notna(val):
            day_num = extract_day_from_date(val)
            month_num = extract_month_from_date(val)
            year_num = extract_year_from_date(val)
            if month_num:
                months_found.add(month_num)
            if year_num:
                years_found.add(year_num)
            if day_num and 1 <= day_num <= 31:
                day_cols.append((i, day_num))
    
    print(f"   📅 Найдено {len(day_cols)} дней, месяц: {months_found}, год: {years_found}")
    
    result = {day: {} for _, day in day_cols}
    doctors_found = 0
    
    for i, row in df.iterrows():
        if len(row) == 0 or pd.isna(row[0]) or str(row[0]).strip() == '':
            continue
        first_cell = str(row[0]).strip()
        
        if first_cell in ['Дата', 'Дежурный врач', 'Сб', 'Вс', 'Пн', 'Вт', 'Ср', 'Чт', 'Пт']:
            continue
        if not first_cell or first_cell == '':
            continue
        
        if not is_valid_doctor_name(first_cell):
            continue
        
        doctor = first_cell
        doctors_found += 1
        for col_idx, day_num in day_cols:
            if col_idx < len(row) and pd.notna(row[col_idx]):
                val = str(row[col_idx]).strip()
                if val and val not in ['', 'nan', 'None']:
                    if re.search(r'\d+[-/]\d+|\d+:\d+', val):
                        normalized = normalize_time(val)
                        result[day_num][(doctor, dept_name)] = normalized
    
    print(f"   👨‍⚕️ Найдено {doctors_found} врачей")
    
    return result, dept_name, months_found, years_found

def parse_oar_graph(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name}...")
    
    df = read_excel_or_csv(file_path)
    if df is None:
        print(f"   ⚠️ Не удалось прочитать файл")
        return {}, dept_name, None, None
    
    header_row_idx = None
    months_found = set()
    years_found = set()
    
    # ============================================================
    # ПОИСК МЕСЯЦА В ТЕКСТЕ (для ОАР)
    # ============================================================
    month_names_ru = {
        'январь': 1, 'февраль': 2, 'март': 3, 'апрель': 4,
        'май': 5, 'июнь': 6, 'июль': 7, 'август': 8,
        'сентябрь': 9, 'октябрь': 10, 'ноябрь': 11, 'декабрь': 12
    }
    
    for i in range(min(5, len(df))):
        row = df.iloc[i].values
        row_str = ' '.join(str(v) for v in row if pd.notna(v)).lower()
        for month_name, month_num in month_names_ru.items():
            if month_name in row_str:
                months_found.add(month_num)
                print(f"   📅 Найден месяц в тексте: {month_name} ({month_num})")
    
    # Ищем строку с Ф.И.О.
    for i, row in df.iterrows():
        row_str = ' '.join(str(v) for v in row.values if pd.notna(v))
        if 'Ф.И.О.' in row_str:
            header_row_idx = i
            break
    
    if header_row_idx is None:
        print(f"   ⚠️ Не найдена строка 'Ф.И.О.'")
        return {}, dept_name, months_found if months_found else None, years_found if years_found else None
    
    header_row = df.iloc[header_row_idx].values
    day_cols = []
    
    for i, val in enumerate(header_row):
        if pd.notna(val):
            day_num = extract_day_from_date(val)
            month_num = extract_month_from_date(val)
            year_num = extract_year_from_date(val)
            if month_num:
                months_found.add(month_num)
            if year_num:
                years_found.add(year_num)
            if day_num and 1 <= day_num <= 31:
                day_cols.append((i, day_num))
    
    print(f"   📅 Найдено {len(day_cols)} дней, месяц: {months_found}, год: {years_found}")
    
    result = {day: {} for _, day in day_cols}
    doctors_found = 0
    
    for i in range(header_row_idx + 1, len(df)):
        row = df.iloc[i].values
        if len(row) == 0 or pd.isna(row[0]) or str(row[0]).strip() == '':
            continue
        doctor = str(row[0]).strip()
        
        if not is_valid_doctor_name(doctor):
            continue
        
        doctors_found += 1
        for col_idx, day_num in day_cols:
            if col_idx < len(row) and pd.notna(row[col_idx]):
                val = str(row[col_idx]).strip()
                if val:
                    match = re.search(r'([АР])', val)
                    if match:
                        result[day_num][(doctor, dept_name)] = match.group(1)
    
    print(f"   👨‍⚕️ Найдено {doctors_found} врачей")
    
    return result, dept_name, months_found if months_found else None, years_found if years_found else None

def parse_surgery_pdf(file_path, dept_name):
    print(f"\n   🔍 Парсим {Path(file_path).name} (PDF)...")
    
    if not PDF_AVAILABLE:
        print(f"   ⚠️ PDF-библиотека не установлена")
        return {}, dept_name, None, None
    
    result = {}
    months_found = set()
    years_found = set()
    
    # Определяем отделение как "Хирургия"
    dept_name = "Хирургия"
    
    try:
        if PDF_LIB == 'pdfplumber':
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        print(f"   📄 Страница {page_num+1}: {len(text)} символов")
                        res, months, years = _parse_surgery_text(text, dept_name)
                        result.update(res)
                        months_found.update(months)
                        years_found.update(years)
        elif PDF_LIB == 'pypdf':
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    print(f"   📄 Страница {page_num+1}: {len(text)} символов")
                    res, months, years = _parse_surgery_text(text, dept_name)
                    result.update(res)
                    months_found.update(months)
                    years_found.update(years)
    except Exception as e:
        print(f"   ⚠️ Ошибка чтения PDF: {e}")
        return {}, dept_name, None, None
    
    print(f"   📅 Найдено {len(result)} дней с дежурствами, месяц: {months_found}, год: {years_found}")
    return result, dept_name, months_found, years_found

def _parse_surgery_text(text, dept_name):
    """Парсит текст из PDF (корректная привязка дат)"""
    result = {}
    months_found = set()
    years_found = set()
    lines = text.split('\n')
    
    date_pattern = re.compile(r'(\d{2})/(\d{2})/(\d{4})')
    time_pattern = re.compile(r'(\d{2}):(\d{2})\s*[-—]\s*(\d{2}):(\d{2})')
    
    # Собираем все строки с данными
    parsed_entries = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if 'Утверждаю' in line or 'Хирургическое' in line or 'График' in line:
            continue
        
        date_match = date_pattern.search(line)
        if date_match:
            day = int(date_match.group(1))
            month = int(date_match.group(2))
            year = int(date_match.group(3))
            months_found.add(month)
            years_found.add(year)
            # Удаляем дату из строки
            line = date_pattern.sub('', line).strip()
            
            # Сохраняем день как "активный"
            current_day = day
        else:
            # Используем последний известный день
            if 'current_day' not in locals():
                continue
        
        # Если строка пустая после удаления даты — пропускаем
        if not line:
            continue
        
        # Ищем время
        time_match = time_pattern.search(line)
        if time_match:
            time_str = f"{time_match.group(1)}/{time_match.group(3)}"
            normalized = normalize_time(time_str)
            doctor_part = line[:time_match.start()].strip()
            doctor = ' '.join(doctor_part.split())
            if doctor and is_valid_doctor_name(doctor):
                parsed_entries.append((current_day, doctor, normalized))
        else:
            # Пробуем найти время в альтернативном формате
            alt_time_match = re.search(r'(\d{2}:\d{2})\s*[-—]\s*(\d{2}:\d{2})', line)
            if alt_time_match:
                time_str = f"{alt_time_match.group(1)[:2]}/{alt_time_match.group(2)[:2]}"
                normalized = normalize_time(time_str)
                doctor_part = line[:alt_time_match.start()].strip()
                doctor = ' '.join(doctor_part.split())
                if doctor and is_valid_doctor_name(doctor):
                    parsed_entries.append((current_day, doctor, normalized))
    
    # ============================================================
    # КОРРЕКТНАЯ ПРИВЯЗКА ДАТ (исправление)
    # ============================================================
    # Проблема: в PDF строки идут так:
    #   день 1: Закороев
    #   день 2: Долгалёв
    #   день 2: Рихан
    #   день 3: Скаковский
    #   
    # Но парсер видит:
    #   день 1: Закороев
    #   день 1: Долгалёв   ← ОШИБКА
    #   день 2: Рихан
    #   день 3: Скаковский
    #
    # Исправление: для каждого дня берём ТОЛЬКО ПЕРВОГО врача,
    # а остальных привязываем к следующему дню.
    
    result = {}
    day_doctors = {}
    
    # Группируем по дням
    for day, doctor, time_str in parsed_entries:
        if day not in day_doctors:
            day_doctors[day] = []
        day_doctors[day].append((doctor, time_str))
    
    # Теперь для каждого дня берём первого врача,
    # а остальных переносим на следующий день
    sorted_days = sorted(day_doctors.keys())
    
    for i, day in enumerate(sorted_days):
        doctors = day_doctors[day]
        
        # Первый врач остаётся в текущем дне
        first_doctor = doctors[0]
        if day not in result:
            result[day] = {}
        result[day][(first_doctor[0], dept_name)] = first_doctor[1]
        
        # Остальные врачи переносятся на следующий день
        if len(doctors) > 1 and i + 1 < len(sorted_days):
            next_day = sorted_days[i + 1]
            for doctor, time_str in doctors[1:]:
                if next_day not in result:
                    result[next_day] = {}
                result[next_day][(doctor, dept_name)] = time_str
        elif len(doctors) > 1 and i + 1 == len(sorted_days):
            # Если это последний день, оставляем всех врачей в нём
            for doctor, time_str in doctors[1:]:
                if day not in result:
                    result[day] = {}
                result[day][(doctor, dept_name)] = time_str
    
    return result, months_found, years_found

# ============================================================
# СОХРАНЕНИЕ В WORD
# ============================================================

def save_to_word(all_data, doctors_by_dept, sorted_depts, output_file, 
                 rukovoditel_text, ploshadka_text, month_num, year, month_name):
    """Сохраняет сводный график в Word с тремя колонками и шапкой"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime
    import calendar
    
    # Создаём документ
    doc = Document()
    
    # Настройка полей (узкие)
    section = doc.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    
    # ============================================================
    # ШАПКА (выравнивание по правому краю)
    # ============================================================
    ruk_parts = rukovoditel_text.split('|')
    ruk_dolzhnost = ruk_parts[0].strip() if len(ruk_parts) > 0 else ''
    ruk_fio = ruk_parts[1].strip() if len(ruk_parts) > 1 else ''
    
    # УТВЕРЖДАЮ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("УТВЕРЖДАЮ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Должность руководителя
    if ruk_dolzhnost:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(ruk_dolzhnost)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    
    # ФИО руководителя (с подчёркиванием)
    if ruk_fio:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(f"____________ {ruk_fio}")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    
    # Дата (сегодняшняя, в родительном падеже)
    today = datetime.datetime.now()
    month_names_ru = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    month_ru = month_names_ru.get(today.month, 'августа')
    date_str = f"{today.day} {month_ru} {today.year} г."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(date_str)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    
    # Пустая строка после шапки (одна)
    doc.add_paragraph()
    
    # ============================================================
    # ЗАГОЛОВОК (выравнивание по центру)
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("ГРАФИК ДЕЖУРСТВА ВРАЧЕЙ СТАЦИОНАРА ГОБУЗ МОКМЦ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(ploshadka_text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    month_names_upper = {
        1: 'ЯНВАРЬ', 2: 'ФЕВРАЛЬ', 3: 'МАРТ', 4: 'АПРЕЛЬ',
        5: 'МАЙ', 6: 'ИЮНЬ', 7: 'ИЮЛЬ', 8: 'АВГУСТ',
        9: 'СЕНТЯБРЬ', 10: 'ОКТЯБРЬ', 11: 'НОЯБРЬ', 12: 'ДЕКАБРЬ'
    }
    month_upper = month_names_upper.get(month_num, 'АВГУСТ')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"НА {month_upper} {year} Г.")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Пустая строка после заголовка (одна)
    doc.add_paragraph()
    
    # ============================================================
    # ДНИ НЕДЕЛИ
    # ============================================================
    _, last_day = calendar.monthrange(year, month_num)
    weekday_names = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
    weekdays = []
    for day in range(1, last_day + 1):
        date_obj = datetime.date(year, month_num, day)
        weekdays.append(weekday_names[date_obj.weekday()])
    
    # ============================================================
    # СБОРКА БЛОКОВ
    # ============================================================
    all_blocks = []
    for day in sorted(all_data.keys()):
        if 1 <= day <= last_day:
            day_name = weekdays[day-1]
            block_lines = [f"{day:02d}.{month_num:02d} {day_name}"]
            
            for dept in sorted_depts:
                dept_doctors = sorted(doctors_by_dept.get(dept, []), key=lambda x: x[0])
                abbr = get_dept_abbr(dept)
                
                for doctor, dept_key in dept_doctors:
                    if (doctor, dept_key) in all_data[day]:
                        value = all_data[day][(doctor, dept_key)]
                        block_lines.append(f"{doctor} ({abbr}) {value}")
            
            all_blocks.append(block_lines)
    
    # ============================================================
    # ФУНКЦИИ ДЛЯ РАСПРЕДЕЛЕНИЯ
    # ============================================================
    def calc_block_height_pt(block, is_last=False):
        line_height_pt = 12
        header_height_pt = 13
        empty_line_pt = 12
        
        total = 0
        for i, line in enumerate(block):
            if i == 0:
                total += header_height_pt
            else:
                total += line_height_pt
        if not is_last:
            total += empty_line_pt
        
        return total
    
    def calc_rows_for_col(blocks):
        total = 0
        for block_idx, block in enumerate(blocks):
            is_last = (block_idx == len(blocks) - 1)
            total += len(block)
            if not is_last:
                total += 1
        return total
    
    # ============================================================
    # НОВЫЙ АЛГОРИТМ РАСПРЕДЕЛЕНИЯ С ОГРАНИЧЕНИЯМИ
    # ============================================================
    FIRST_PAGE_MAX_ROWS = 48
    OTHER_PAGE_MAX_ROWS = 59
    
    def distribute_blocks_with_limit(blocks, first_max, other_max):
        """Распределяет блоки по страницам с учётом ограничений"""
        pages = []
        current_page = [[], [], []]
        col_heights = [0, 0, 0]
        current_col = 0
        is_first_page = True
        
        for block in blocks:
            block_height = calc_block_height_pt(block, is_last=False)
            # Пересчитываем в строки (примерно)
            block_rows = len(block) + 1  # +1 для пустой строки
            
            # Определяем максимальное количество строк для текущей страницы
            if is_first_page:
                max_rows = first_max
            else:
                max_rows = other_max
            
            # Проверяем, помещается ли блок в текущую колонку
            if col_heights[current_col] + block_rows <= max_rows:
                current_page[current_col].append(block)
                col_heights[current_col] += block_rows
            else:
                # Ищем следующую колонку
                col_found = False
                for next_col in range(current_col + 1, 3):
                    if col_heights[next_col] + block_rows <= max_rows:
                        current_col = next_col
                        current_page[current_col].append(block)
                        col_heights[current_col] += block_rows
                        col_found = True
                        break
                
                if not col_found:
                    # Блок не поместился — сохраняем страницу и создаём новую
                    if any(len(col) > 0 for col in current_page):
                        pages.append(current_page)
                    
                    # Создаём новую страницу
                    current_page = [[], [], []]
                    col_heights = [0, 0, 0]
                    current_col = 0
                    is_first_page = False
                    
                    # Добавляем блок в новую страницу
                    current_page[current_col].append(block)
                    col_heights[current_col] += block_rows
        
        # Добавляем последнюю страницу
        if any(len(col) > 0 for col in current_page):
            pages.append(current_page)
        
        return pages
    
    # ============================================================
    # РАСПРЕДЕЛЯЕМ БЛОКИ
    # ============================================================
    pages = distribute_blocks_with_limit(all_blocks, FIRST_PAGE_MAX_ROWS, OTHER_PAGE_MAX_ROWS)
    
    if not pages:
        pages = [[[], [], []]]
    
    # ============================================================
    # ФУНКЦИЯ ДЛЯ ЗАПОЛНЕНИЯ КОЛОНКИ
    # ============================================================
    def fill_column(table, col_idx, blocks, start_row, max_rows):
        row = start_row
        for block_idx, block in enumerate(blocks):
            if row >= max_rows:
                break
                
            cell = table.cell(row, col_idx)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            run = p.add_run(block[0])
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            row += 1
            
            for line in block[1:]:
                if row >= max_rows:
                    break
                cell = table.cell(row, col_idx)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                run = p.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                row += 1
            
            is_last_block = (block_idx == len(blocks) - 1)
            if not is_last_block and row < max_rows:
                cell = table.cell(row, col_idx)
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.add_run('')
                row += 1
        
        return row
    
    # ============================================================
    # СОЗДАНИЕ ТАБЛИЦ
    # ============================================================
    first_page = True
    for page_idx, page in enumerate(pages):
        if not any(len(col) > 0 for col in page):
            continue
        
        if not first_page:
            doc.add_page_break()
        
        first_page = False
        
        max_rows = max(
            calc_rows_for_col(page[0]),
            calc_rows_for_col(page[1]),
            calc_rows_for_col(page[2])
        )
        table_height = max(max_rows, 3)
        
        table = doc.add_table(rows=table_height, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        fill_column(table, 0, page[0], 0, table_height)
        fill_column(table, 1, page[1], 0, table_height)
        fill_column(table, 2, page[2], 0, table_height)
        
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(2.5)
    
    doc.save(output_file)
    print(f"💾 Word сохранён: {output_file}")

# ============================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================
def build_master_schedule(input_folder, rukovoditel_text, ploshadka_text, 
                          selected_month, selected_year):
    """Собирает сводный график из файлов в папке"""
    print(f"\n📊 Начинаем сборку сводного графика...")
    print(f"📁 Папка: {input_folder}")
    print(f"📅 Месяц: {selected_month}, Год: {selected_year}")
    print("=" * 60)
    
    folder = Path(input_folder)
    
    # Создаём папку для результата
    output_folder = folder / "Сводный график"
    output_folder.mkdir(exist_ok=True)
    
    # Ищем все файлы
    all_files = []
    for ext in ['*.xlsx', '*.xls', '*.csv', '*.pdf', '*.txt']:
        all_files.extend(folder.glob(ext))
    
    all_files = [f for f in all_files if not f.name.startswith('~')]
    all_files = [f for f in all_files if 'сводный' not in f.name.lower()]
    
    if not all_files:
        print("❌ Не найдено ни одного файла с графиками!")
        return False
    
    print(f"\n📄 Найдено {len(all_files)} файлов:")
    for f in all_files:
        print(f"   - {f.name}")
    
    print("\n" + "=" * 60)
    
    all_data = {}
    doctors_by_dept = {}
    dept_order_map = {}
    files_with_errors = []
    
    for file_path in all_files:
        print(f"\n{'='*60}")
        print(f"📂 Обрабатываем: {file_path.name}")
        
        df_temp = read_excel_or_csv(file_path)
        if df_temp is not None:
            dept_name = detect_department(df_temp, file_path)
        else:
            dept_name = Path(file_path).stem.replace('_', ' ')
        
        print(f"   🏥 Отделение: {dept_name}")
        print(f"   📛 Сокращение: {get_dept_abbr(dept_name)}")
        
        if dept_name == 'ОАР' or 'ОАР' in file_path.name:
            data, detected_dept, months, years = parse_oar_graph(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        elif dept_name == 'Ответственные по стационару' or 'ответственн' in dept_name.lower() or 'Ответственные' in file_path.name:
            data, detected_dept, months, years = parse_responsible_graph(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        elif file_path.suffix.lower() == '.pdf':
            data, detected_dept, months, years = parse_surgery_pdf(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        elif file_path.suffix.lower() == '.txt':
            data, detected_dept, months, years = parse_surgery_pdf(file_path, dept_name)
            if not data:
                print(f"   ⚠️ Не удалось прочитать TXT, пропускаем")
                continue
        else:
            data, detected_dept, months, years = parse_standard_graph(file_path, dept_name)
            if detected_dept and detected_dept != dept_name:
                dept_name = detected_dept
        
        if not data:
            print(f"   ⚠️ Данные не найдены для {dept_name}")
            continue
        
        # Проверяем месяц и год
        if months:
            if selected_month not in months:
                print(f"   ⚠️ Месяц в файле ({months}) не соответствует выбранному ({selected_month})")
                files_with_errors.append((file_path.name, f"Месяц: {months}, ожидался: {selected_month}"))
                continue
        
        if years:
            if selected_year not in years:
                print(f"   ⚠️ Год в файле ({years}) не соответствует выбранному ({selected_year})")
                files_with_errors.append((file_path.name, f"Год: {years}, ожидался: {selected_year}"))
                continue
        
        for day, doctors in data.items():
            if day not in all_data:
                all_data[day] = {}
            all_data[day].update(doctors)
            
            if dept_name not in doctors_by_dept:
                doctors_by_dept[dept_name] = []
            for doctor_key in doctors.keys():
                if doctor_key not in doctors_by_dept[dept_name]:
                    doctors_by_dept[dept_name].append(doctor_key)
        
        dept_order_map[dept_name] = get_dept_order(dept_name)
        
        print(f"   ✅ Обработано {len(data)} дней, {len(doctors_by_dept.get(dept_name, []))} врачей")
    
    if not all_data:
        print("❌ Данные не найдены ни в одном файле!")
        # ============================================================
        # ПОКАЗЫВАЕМ ПОНЯТНОЕ СООБЩЕНИЕ В GUI (ДИНАМИЧЕСКОЕ)
        # ============================================================
        try:
            root = tk._default_root
            if root:
                month_name_ru = MONTHS_RU_NOMINATIVE.get(selected_month, '')
                msg = f"Данные за {month_name_ru} {selected_year} года не найдены ни в одном из графиков.\n\n"
                if files_with_errors:
                    msg += "Проверьте, что в файлах указан правильный месяц и год.\n\n"
                    msg += "Исключённые файлы:\n"
                    for fname, error in files_with_errors[:5]:
                        msg += f"  • {fname}\n"
                    if len(files_with_errors) > 5:
                        msg += f"  ... и ещё {len(files_with_errors) - 5} файлов\n"
                else:
                    msg += "Возможно, в папке нет файлов графиков за выбранный период.\n"
                    msg += "Проверьте выбранную папку и месяц."
                messagebox.showerror("Данные не найдены", msg)
        except:
            pass
        return False

    # ============================================================
    # ПОКАЗЫВАЕМ СПИСОК ИСКЛЮЧЁННЫХ ФАЙЛОВ (ЕСЛИ ЕСТЬ)
    # ============================================================
    if files_with_errors:
        print("\n" + "=" * 60)
        print("⚠️⚠️⚠️ Файлы, исключённые из сводного графика (не соответствуют выбранному месяцу/году):")
        print("-" * 60)
        for fname, error in files_with_errors:
            print(f"   ❌ {fname}")
            print(f"      {error}")
        print("-" * 60)
        print(f"   Всего исключено: {len(files_with_errors)} файлов")
        print("=" * 60)
        
        # ПОКАЗЫВАЕМ В GUI (если есть данные)
        try:
            root = tk._default_root
            if root and all_data:
                msg = "Следующие файлы не соответствуют выбранному месяцу/году и были исключены:\n\n"
                for fname, error in files_with_errors[:5]:
                    msg += f"• {fname}\n  ({error})\n"
                if len(files_with_errors) > 5:
                    msg += f"\n... и ещё {len(files_with_errors) - 5} файлов"
                msg += f"\n\nВсего исключено: {len(files_with_errors)} файлов"
                messagebox.showwarning("Файлы исключены", msg)
        except:
            pass

    # ============================================================
    # СОРТИРОВКА ОТДЕЛЕНИЙ
    # ============================================================
    sorted_depts = sorted(
        doctors_by_dept.keys(),
        key=lambda d: (dept_order_map.get(d, 999), d)
    )

    print("\n" + "=" * 60)
    print(f"✅ Всего найдено {len(all_data)} дней с дежурствами")
    print(f"📊 Найдено отделений: {len(doctors_by_dept)}")
    for dept in sorted_depts:
        abbr = get_dept_abbr(dept)
        print(f"   {dept} ({abbr}): {len(doctors_by_dept[dept])} врачей")

    month_name = MONTHS_RU_UPPER.get(selected_month, 'АВГУСТ')
    
    # Сохраняем TXT
    _, last_day = calendar.monthrange(selected_year, selected_month)
    weekday_names = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
    weekdays = []
    for day in range(1, last_day + 1):
        date_obj = datetime.date(selected_year, selected_month, day)
        weekdays.append(weekday_names[date_obj.weekday()])
    
    output_lines = []
    for day in sorted(all_data.keys()):
        if 1 <= day <= last_day:
            day_name = weekdays[day-1]
            output_lines.append(f"{day:02d}.{selected_month:02d} {day_name}")
            
            for dept in sorted_depts:
                dept_doctors = sorted(doctors_by_dept.get(dept, []), key=lambda x: x[0])
                abbr = get_dept_abbr(dept)
                
                for doctor, dept_key in dept_doctors:
                    if (doctor, dept_key) in all_data[day]:
                        value = all_data[day][(doctor, dept_key)]
                        output_lines.append(f"{doctor} ({abbr}) {value}")
            
            output_lines.append("")
    
    month_lower = MONTHS_RU.get(selected_month, 'августа')
    txt_file = output_folder / f"сводный_график_{month_lower}_{selected_year}.txt"
    with open(txt_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))
    
    print(f"\n💾 TXT сохранён: {txt_file}")
    
    # Сохраняем Word
    docx_file = output_folder / f"сводный_график_{month_lower}_{selected_year}.docx"
    save_to_word(all_data, doctors_by_dept, sorted_depts, docx_file,
                 rukovoditel_text, ploshadka_text, selected_month, selected_year, month_name)
    
    print(f"\n✅ Готово!")
    print(f"   📄 TXT: {txt_file}")
    print(f"   📝 Word: {docx_file}")
    
    return True

# ============================================================
# ГРАФИЧЕСКИЙ ИНТЕРФЕЙС
# ============================================================
class App:
    def __init__(self, root):
        self.root = root
        root.title("Сводный график дежурств")
        root.geometry("650x680")
        root.resizable(False, False)
        
        # Переменные
        self.folder_path = tk.StringVar(value="")
        self.rukovoditel = tk.StringVar()
        self.ploshadka = tk.StringVar()
        self.month_var = tk.StringVar()
        self.year_var = tk.StringVar()
        
        # Загружаем справочники
        self.rukovoditeli_list = load_spravochnik('spravochnik_rukovoditeli.txt')
        self.ploshadki_list = load_spravochnik('spravochnik_ploshadki.txt')
        
        # Определяем год по умолчанию
        now = datetime.datetime.now()
        default_year = now.year
        if now.month == 12 and now.day > 15:
            default_year += 1
        
        self.year_var.set(str(default_year))
        
        # Создаём интерфейс
        self.create_widgets()
        
    def create_widgets(self):
        # Заголовок
        tk.Label(self.root, text="Сводный график дежурств", 
                font=("Arial", 16, "bold")).pack(pady=10)
        
        # Рамка для выбора папки
        frame_folder = tk.LabelFrame(self.root, text="Папка с графиками", padx=10, pady=10)
        frame_folder.pack(fill="x", padx=20, pady=5)
        
        tk.Label(frame_folder, text="Выберите папку, содержащую файлы графиков отделений:").pack(anchor="w")
        
        folder_row = tk.Frame(frame_folder)
        folder_row.pack(fill="x", pady=5)
        tk.Entry(folder_row, textvariable=self.folder_path, width=50).pack(side="left", padx=(0,5))
        tk.Button(folder_row, text="Обзор...", command=self.select_folder).pack(side="left")
        
        # Рамка для месяца и года
        frame_date = tk.LabelFrame(self.root, text="Период", padx=10, pady=10)
        frame_date.pack(fill="x", padx=20, pady=5)
        
        date_row = tk.Frame(frame_date)
        date_row.pack(fill="x", pady=5)
        
        tk.Label(date_row, text="Месяц:").pack(side="left", padx=(0,10))
        months_list = [MONTHS_RU_NOMINATIVE[i] for i in range(1, 13)]
        combo_month = ttk.Combobox(date_row, textvariable=self.month_var, 
                                values=months_list, width=20)
        combo_month.pack(side="left", padx=(0,20))
        self.month_var.set("Август")
        
        tk.Label(date_row, text="Год:").pack(side="left", padx=(0,10))
        years_list = [str(y) for y in range(2024, 2031)]
        combo_year = ttk.Combobox(date_row, textvariable=self.year_var, 
                                values=years_list, width=10)
        combo_year.pack(side="left")
        
        # Рамка для руководителя
        frame_ruk = tk.LabelFrame(self.root, text="Руководитель", padx=10, pady=10)
        frame_ruk.pack(fill="x", padx=20, pady=5)
        
        if self.rukovoditeli_list:
            tk.Label(frame_ruk, text="Выберите руководителя:").pack(anchor="w")
            combo_ruk = ttk.Combobox(frame_ruk, textvariable=self.rukovoditel, 
                                    values=self.rukovoditeli_list, width=80)
            combo_ruk.pack(anchor="w", pady=5)
            if self.rukovoditeli_list:
                self.rukovoditel.set(self.rukovoditeli_list[1] if len(self.rukovoditeli_list) > 1 else self.rukovoditeli_list[0])
        else:
            tk.Label(frame_ruk, text="Нет данных в справочнике руководителей").pack(anchor="w")
        
        # Рамка для площадки
        frame_plo = tk.LabelFrame(self.root, text="Площадка", padx=10, pady=10)
        frame_plo.pack(fill="x", padx=20, pady=5)
        
        if self.ploshadki_list:
            tk.Label(frame_plo, text="Выберите площадку:").pack(anchor="w")
            combo_plo = ttk.Combobox(frame_plo, textvariable=self.ploshadka, 
                                    values=self.ploshadki_list, width=60)
            combo_plo.pack(anchor="w", pady=5)
            if self.ploshadki_list:
                self.ploshadka.set(self.ploshadki_list[0])
        else:
            tk.Label(frame_plo, text="Нет данных в справочнике площадок").pack(anchor="w")
        
        # Рамка для списка файлов
        frame_files = tk.LabelFrame(self.root, text="Файлы в папке", padx=10, pady=10)
        frame_files.pack(fill="both", expand=True, padx=20, pady=5)
        
        self.files_listbox = tk.Listbox(frame_files, height=6)
        self.files_listbox.pack(fill="both", expand=True)
        
        btn_refresh = tk.Button(frame_files, text="🔄 Обновить список", 
                                command=self.refresh_files_list)
        btn_refresh.pack(pady=5)
        
        # ============================================================
        # КНОПКА "СОСТАВИТЬ СВОДНЫЙ ГРАФИК"
        # ============================================================
        btn_run = tk.Button(self.root, text="Составить сводный график", 
                           command=self.run, bg="#4CAF50", fg="white",
                           font=("Arial", 12, "bold"), padx=20, pady=10)
        btn_run.pack(pady=20)
        
        # Статус
        self.status_label = tk.Label(self.root, text="", fg="blue")
        self.status_label.pack(pady=5)

    def refresh_files_list(self):
        """Обновляет список файлов в папке"""
        folder = self.folder_path.get().strip()
        self.files_listbox.delete(0, tk.END)
        
        if not folder:
            self.files_listbox.insert(tk.END, "⚠️ Папка не выбрана")
            return
        
        if not os.path.exists(folder):
            self.files_listbox.insert(tk.END, "⚠️ Папка не существует")
            return
        
        try:
            files = os.listdir(folder)
            graph_files = sorted([f for f in files if f.endswith(('.xlsx', '.xls', '.csv', '.pdf', '.txt'))])
            if graph_files:
                for f in graph_files:
                    self.files_listbox.insert(tk.END, f)
                self.status_label.config(text=f"📁 Найдено {len(graph_files)} файлов", fg="green")
            else:
                self.files_listbox.insert(tk.END, "⚠️ Нет файлов графиков")
                self.status_label.config(text="⚠️ В папке нет файлов графиков", fg="orange")
        except Exception as e:
            self.files_listbox.insert(tk.END, f"❌ Ошибка: {str(e)}")
            self.status_label.config(text=f"❌ Ошибка: {str(e)}", fg="red")

    def select_folder(self):
        folder = filedialog.askdirectory(title="Выберите папку с графиками")
        if folder:
            self.folder_path.set(folder)
            self.refresh_files_list()

    def run(self):
        folder = self.folder_path.get().strip()
        if not folder:
            messagebox.showerror("Ошибка", "Выберите папку с графиками!")
            return
        
        if not os.path.exists(folder):
            messagebox.showerror("Ошибка", f"Папка не существует: {folder}")
            return
        
        ruk_text = self.rukovoditel.get().strip()
        if not ruk_text:
            messagebox.showerror("Ошибка", "Выберите руководителя!")
            return
        
        plo_text = self.ploshadka.get().strip()
        if not plo_text:
            messagebox.showerror("Ошибка", "Выберите площадку!")
            return
        
        # Парсим месяц (с поддержкой именительного падежа из списка)
        month_str = self.month_var.get().strip()
        
        # Сначала пробуем найти в словаре именительного падежа (для отображения в списке)
        month_names_nominative = {v: k for k, v in MONTHS_RU_NOMINATIVE.items()}
        if month_str in month_names_nominative:
            month_num = month_names_nominative[month_str]
        else:
            # Если не нашли, пробуем в родительном падеже
            month_str_lower = month_str.lower()
            month_names = {v.lower(): k for k, v in MONTHS_RU.items()}
            if month_str_lower in month_names:
                month_num = month_names[month_str_lower]
            else:
                messagebox.showerror("Ошибка", f"Неверный месяц: {month_str}")
                return
        
        # Парсим год
        try:
            year_num = int(self.year_var.get().strip())
        except:
            messagebox.showerror("Ошибка", "Неверный формат года!")
            return
        
        self.status_label.config(text="⏳ Обработка...")
        self.root.update()
        
        try:
            # Перехватываем вывод в консоль для показа в GUI
            import io
            from contextlib import redirect_stdout
            
            # Сохраняем оригинальный stdout
            original_stdout = sys.stdout
            # Создаём буфер для захвата вывода
            captured_output = io.StringIO()
            
            # Запускаем с захватом вывода
            with redirect_stdout(captured_output):
                success = build_master_schedule(folder, ruk_text, plo_text, month_num, year_num)
            
            # Возвращаем stdout
            sys.stdout = original_stdout
            
            # Показываем захваченный вывод
            output_text = captured_output.getvalue()
            if output_text:
                print(output_text)  # Выводим в консоль для отладки
            
            if success:
                self.status_label.config(text="✅ Готово! Файлы созданы в папке 'Сводный график'", fg="green")
                messagebox.showinfo("Успех", "Сводный график успешно создан!\nФайлы сохранены в папке 'Сводный график'")
            else:
                self.status_label.config(text="❌ Ошибка при создании графика", fg="red")
                messagebox.showerror("Ошибка", "Не удалось создать сводный график.\nПроверьте консоль для подробностей.")
        except Exception as e:
            self.status_label.config(text=f"❌ Ошибка: {str(e)}", fg="red")
            messagebox.showerror("Ошибка", str(e))

# ============================================================
# ЗАПУСК
# ============================================================
if __name__ == '__main__':
    root = tk.Tk()
    app = App(root)
    root.mainloop()